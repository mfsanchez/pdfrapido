#!/usr/bin/env python3
"""
Convierte PDF con texto (incluido OCR) a DOCX limpio y legible.
Detecta PDFs OCR (bloques de <5 palabras) y reagrupa tokens por posición Y.
Uso: python3 pdftext_to_docx.py input.pdf output.docx
"""
import sys, re
import fitz
from docx import Document
from docx.shared import Pt

LINE_MERGE_GAP  = 5   # pt — mismo renglón si distancia Y < 5
PARA_BREAK_GAP  = 12  # pt — párrafo nuevo si salto vertical > 12
OCR_WORD_THRESH = 5   # media de palabras por bloque; si < 5 → PDF OCR

def is_ocr_pdf(doc):
    counts = []
    for page in doc:
        for block in page.get_text("blocks"):
            if block[6] == 0:
                words = block[4].split()
                if words:
                    counts.append(len(words))
    if not counts:
        return False
    return (sum(counts) / len(counts)) < OCR_WORD_THRESH

def extract_clean_paragraphs(doc):
    paragraphs = []

    for page in doc:
        spans = []
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span["text"].strip()
                    if text:
                        y = span["origin"][1]
                        x = span["bbox"][0]
                        spans.append((y, x, text))

        if not spans:
            continue

        spans.sort(key=lambda s: (round(s[0] / LINE_MERGE_GAP), s[1]))

        # Fusionar spans en líneas reales por proximidad Y
        lines = []
        cur_y = spans[0][0]
        cur_tokens = [spans[0][2]]

        for y, x, text in spans[1:]:
            if abs(y - cur_y) <= LINE_MERGE_GAP:
                cur_tokens.append(text)
            else:
                lines.append((cur_y, ' '.join(cur_tokens)))
                cur_y = y
                cur_tokens = [text]
        lines.append((cur_y, ' '.join(cur_tokens)))

        # Agrupar líneas en párrafos por salto vertical
        para_lines = [lines[0][1]]
        prev_y = lines[0][0]

        for y, text in lines[1:]:
            if y - prev_y > PARA_BREAK_GAP:
                combined = ' '.join(para_lines).strip()
                if combined:
                    paragraphs.append(combined)
                para_lines = [text]
            else:
                para_lines.append(text)
            prev_y = y

        combined = ' '.join(para_lines).strip()
        if combined:
            paragraphs.append(combined)

    return paragraphs

def sanitize(s):
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', s)

HEADING_KW = re.compile(
    r'^(art[íi]culo|art\.|cl[áa]usula|secci[óo]n|cap[íi]tulo|t[íi]tulo|anexo|'
    r'apartado|disposici[óo]n|p[áa]rrafo)\b',
    re.IGNORECASE
)

def classify(para):
    s = para.strip()
    upper_ratio = sum(1 for c in s if c.isupper()) / max(len(s), 1)
    if len(s) <= 80 and upper_ratio > 0.6 and len(s) > 3:
        return 'heading1'
    # Encabezado legal: "Cláusula Primera...", "Artículo 5...", "Capítulo I..."
    if HEADING_KW.match(s) and len(s) < 120:
        return 'heading2'
    # Encabezado numerado: "1. Objeto", "2) Ámbito"
    if re.match(r'^\d+[\.\)]\s+\S', s) and len(s) < 100:
        return 'heading2'
    return 'body'

def main():
    if len(sys.argv) < 3:
        print("Uso: pdftext_to_docx.py input.pdf output.docx", file=sys.stderr)
        sys.exit(1)

    input_path, output_path = sys.argv[1], sys.argv[2]

    pdf = fitz.open(input_path)
    ocr_mode = is_ocr_pdf(pdf)
    paragraphs = extract_clean_paragraphs(pdf)
    pdf.close()

    if not paragraphs or all(len(p.strip()) < 3 for p in paragraphs):
        print("ERROR: no se pudo extraer texto del PDF", file=sys.stderr)
        sys.exit(2)

    doc = Document()
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.name = 'Calibri'

    for para in paragraphs:
        para = sanitize(para.strip())
        if not para:
            continue
        kind = classify(para)
        if kind == 'heading1':
            doc.add_heading(para, level=1)
        elif kind == 'heading2':
            doc.add_heading(para, level=2)
        else:
            doc.add_paragraph(para)

    doc.save(output_path)

    if ocr_mode:
        print("OCR_MODE", file=sys.stderr)
    print(f"OK: {output_path}")

if __name__ == '__main__':
    main()
